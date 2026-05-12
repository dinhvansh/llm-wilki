from __future__ import annotations

import re
import shutil
from html.parser import HTMLParser
from collections import Counter
from dataclasses import asdict, dataclass
from pathlib import Path
from urllib.parse import quote
from uuid import uuid4

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from pypdf import PdfReader

from app.config import settings
from app.core.docling_parser import parse_with_docling
from app.core.llm_client import llm_client
from app.core.reliability import PROMPT_VERSION
from app.core.runtime_config import load_runtime_snapshot


TOKEN_RE = re.compile(r"[a-zA-Z0-9][a-zA-Z0-9_-]+")
SENTENCE_RE = re.compile(r"(?<=[.!?])\s+")
DRAWING_XPATH = ".//w:drawing//a:blip"


@dataclass
class ParsedDocument:
    text: str
    mime_type: str
    source_type: str
    metadata: dict


@dataclass
class IngestStageResult:
    name: str
    status: str
    details: dict


@dataclass
class IngestArtifacts:
    parsed: ParsedDocument
    chunks: list[dict]
    summary: str
    key_facts: list[str]
    tags: list[str]
    entities: list[dict]
    claims: list[dict]
    timeline_events: list[dict]
    glossary_terms: list[dict]
    page_type_candidates: list[dict]
    stage_results: list[IngestStageResult]


def ensure_upload_dir() -> Path:
    upload_dir = Path(settings.UPLOAD_DIR)
    upload_dir.mkdir(parents=True, exist_ok=True)
    return upload_dir


def infer_source_type(filename: str) -> str:
    extension = Path(filename).suffix.lower()
    return {
        ".md": "markdown",
        ".markdown": "markdown",
        ".txt": "txt",
        ".pdf": "pdf",
        ".docx": "docx",
        ".png": "image_ocr",
        ".jpg": "image_ocr",
        ".jpeg": "image_ocr",
        ".webp": "image_ocr",
        ".tif": "image_ocr",
        ".tiff": "image_ocr",
    }.get(extension, "txt")


def infer_mime_type(filename: str, provided: str | None) -> str:
    if provided:
        return provided
    extension = Path(filename).suffix.lower()
    return {
        ".md": "text/markdown",
        ".markdown": "text/markdown",
        ".txt": "text/plain",
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".webp": "image/webp",
        ".tif": "image/tiff",
        ".tiff": "image/tiff",
    }.get(extension, "application/octet-stream")


def parse_document(path: Path, mime_type: str, source_type: str) -> ParsedDocument:
    if source_type in {"txt", "transcript", "url"}:
        text = path.read_text(encoding="utf-8", errors="ignore")
        return ParsedDocument(text=normalize_text(text), mime_type=mime_type, source_type=source_type, metadata={})

    if source_type in {"markdown", "pdf", "docx", "image_ocr"}:
        text, metadata = parse_with_docling(path, mime_type, source_type)
        metadata = dict(metadata or {})
        if source_type == "docx":
            try:
                document = Document(path)
                ordered_blocks, image_urls = parse_docx_blocks(document, path)
                metadata = {
                    **metadata,
                    "orderedBlocks": ordered_blocks,
                    "images": image_urls,
                    "imageCount": len(image_urls),
                    "tableCount": sum(1 for block in ordered_blocks if block.get("type") == "table"),
                }
            except Exception:
                # Keep docling output if block-level enrichment fails.
                metadata = metadata
        elif source_type == "image_ocr":
            image_url = public_upload_url(path)
            metadata = {
                **metadata,
                "orderedBlocks": [{"type": "image", "url": image_url, "alt": path.stem or "Uploaded image"}],
                "images": [image_url],
                "imageCount": 1,
            }
        return ParsedDocument(
            text=normalize_text(text),
            mime_type=mime_type,
            source_type=source_type,
            metadata=metadata,
        )

    text = path.read_text(encoding="utf-8", errors="ignore")
    return ParsedDocument(text=normalize_text(text), mime_type=mime_type, source_type=source_type, metadata={})


class ReadableHTMLParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.title_parts: list[str] = []
        self.text_parts: list[str] = []
        self._tag_stack: list[str] = []
        self._skip_depth = 0

    def handle_starttag(self, tag: str, attrs) -> None:
        lowered = tag.lower()
        self._tag_stack.append(lowered)
        if lowered in {"script", "style", "noscript", "svg"}:
            self._skip_depth += 1
        if lowered in {"p", "div", "section", "article", "li", "br", "h1", "h2", "h3", "h4"}:
            self.text_parts.append("\n")

    def handle_endtag(self, tag: str) -> None:
        lowered = tag.lower()
        if lowered in {"script", "style", "noscript", "svg"} and self._skip_depth > 0:
            self._skip_depth -= 1
        if self._tag_stack:
            self._tag_stack.pop()
        if lowered in {"p", "div", "section", "article", "li", "h1", "h2", "h3", "h4"}:
            self.text_parts.append("\n")

    def handle_data(self, data: str) -> None:
        text = " ".join(data.split())
        if not text or self._skip_depth:
            return
        if self._tag_stack and self._tag_stack[-1] == "title":
            self.title_parts.append(text)
            return
        self.text_parts.append(text)

    @property
    def title(self) -> str:
        return normalize_text(" ".join(self.title_parts))

    @property
    def text(self) -> str:
        return normalize_text(" ".join(self.text_parts))


def extract_readable_html(raw_html: str) -> tuple[str, str]:
    parser = ReadableHTMLParser()
    parser.feed(raw_html)
    return parser.title, parser.text


def normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def tokenize(text: str) -> list[str]:
    return [token.lower() for token in TOKEN_RE.findall(text)]


def promote_heading_lines(text: str) -> str:
    lines = text.splitlines()
    promoted: list[str] = []

    def next_non_empty(index: int) -> str:
        for candidate in lines[index + 1 :]:
            stripped = candidate.strip()
            if stripped:
                return stripped
        return ""

    for index, raw_line in enumerate(lines):
        stripped = raw_line.strip()
        if not stripped or stripped.startswith("#") or stripped.startswith("|") or stripped.startswith("```"):
            promoted.append(raw_line)
            continue
        words = stripped.split()
        if (
            1 <= len(words) <= 6
            and len(stripped) <= 60
            and not re.search(r"[.!?;:]$", stripped)
            and stripped == stripped.title()
            and next_non_empty(index)
        ):
            promoted.append(f"## {stripped}")
            continue
        promoted.append(raw_line)
    return "\n".join(promoted)


def detect_language(text: str) -> str:
    sample = (text or "")[:4000].lower()
    if re.search(r"[ăâđêôơư]", sample):
        return "vi"
    common_vi = sum(1 for token in (" va ", " cua ", " cho ", " quy trinh ", " chinh sach ", " duoc ") if token in f" {sample} ")
    common_en = sum(1 for token in (" the ", " and ", " policy ", " process ", " must ", " should ") if token in f" {sample} ")
    return "vi" if common_vi > common_en else "en"


def classify_document_profile(title: str, parsed: ParsedDocument, chunks: list[dict], summary: str = "") -> dict:
    combined = "\n".join([title, summary, parsed.text[:12000]]).lower()
    source_type = (parsed.source_type or "").lower()
    scores = {
        "policy": 0.0,
        "sop": 0.0,
        "meeting_note": 0.0,
        "report": 0.0,
        "reference": 0.0,
        "user_note": 0.0,
    }

    if re.search(r"\bpolicy\b|\bstandard\b|\bgovernance\b|\bcompliance\b|\bshall\b|\bmust\b|\brequired\b", combined):
        scores["policy"] += 0.75
    if re.search(r"\bsop\b|\bprocedure\b|\bworkflow\b|\bchecklist\b|\bstep\b|\bhow to\b", combined):
        scores["sop"] += 0.72
    if re.search(r"\bmeeting\b|\bminutes\b|\battendees\b|\baction items?\b|\bdecisions?\b|\bdiscussion\b", combined):
        scores["meeting_note"] += 0.78
    if re.search(r"\breport\b|\bexecutive summary\b|\bfindings\b|\banalysis\b|\brecommendations?\b|\bcurrent state\b|\bgoal\b", combined):
        scores["report"] += 0.7
    if re.search(r"\breference\b|\bglossary\b|\bapi\b|\bfield\b|\bparameter\b|\bdefinition\b|\bfaq\b", combined):
        scores["reference"] += 0.66
    if re.search(r"\bnotes?\b|\bbrainstorm\b|\bdraft\b|\bidea\b|\bpersonal\b", combined):
        scores["user_note"] += 0.58

    if source_type == "transcript":
        scores["meeting_note"] += 0.28
    if source_type == "url":
        scores["reference"] += 0.12
    if source_type == "txt":
        scores["user_note"] += 0.1

    numbered_steps = sum(1 for chunk in chunks[:10] if re.search(r"(?:^|\s)\d+\.(?:\s|$)|\bstep\s+\d+\b", chunk.get("content", "").lower()))
    if numbered_steps >= 2:
        scores["sop"] += 0.2

    selected_type = max(scores.items(), key=lambda item: item[1])[0]
    confidence = max(scores.values())
    if confidence < 0.45:
        selected_type = "report" if source_type in {"pdf", "docx", "markdown"} else "reference"
        confidence = 0.45

    reasons: list[str] = []
    for label, pattern in [
        ("policy_language", r"\bpolicy\b|\bshall\b|\bmust\b|\brequired\b"),
        ("procedure_markers", r"\bprocedure\b|\bworkflow\b|\bstep\b|\bchecklist\b"),
        ("meeting_markers", r"\bmeeting\b|\bminutes\b|\baction items?\b"),
        ("report_markers", r"\bexecutive summary\b|\bfindings\b|\brecommendations?\b"),
        ("reference_markers", r"\bglossary\b|\bdefinition\b|\bfield\b|\bparameter\b"),
    ]:
        if re.search(pattern, combined):
            reasons.append(label)

    return {
        "documentType": selected_type,
        "confidence": round(min(confidence, 0.95), 2),
        "reasons": reasons[:5],
        "sourceType": source_type,
    }


def annotate_chunk_section_roles(chunks: list[dict], document_type: str) -> list[dict]:
    def infer_role(chunk: dict) -> str:
        heading_path = chunk.get("metadata", {}).get("headingPath") or []
        heading_text = " ".join([str(part) for part in heading_path if part]).lower()
        content_text = str(chunk.get("content") or "").lower()
        section_text = " ".join(filter(None, [str(chunk.get("section_title") or ""), heading_text, content_text[:280]]))

        if document_type == "policy":
            if re.search(r"\bscope\b|\bapplies to\b|\bcoverage\b", section_text):
                return "scope"
            if re.search(r"\bexception\b|\bwaiver\b|\bunless\b", section_text):
                return "exception"
            if re.search(r"\bowner\b|\bresponsib", section_text):
                return "owner"
            if re.search(r"\bmust\b|\bshall\b|\brequired\b|\bprohibited\b|\bnot allowed\b", section_text):
                return "rule"
        if document_type == "sop":
            if re.search(r"\bprerequisite\b|\bbefore you start\b|\binputs?\b", section_text):
                return "prerequisite"
            if re.search(r"\bvalidate\b|\bcheck\b|\bverify\b|\bacceptance\b", section_text):
                return "validation"
            if re.search(r"(?:^|\s)\d+\.(?:\s|$)|\bstep\s+\d+\b|\bprocedure\b", section_text):
                return "step"
        if document_type == "meeting_note":
            if re.search(r"\bdecision\b|\bagreed\b|\bapproved\b", section_text):
                return "decision"
            if re.search(r"\baction items?\b|\bnext steps?\b|\bowner\b", section_text):
                return "action_item"
            if re.search(r"\brisk\b|\bissue\b|\bblocker\b", section_text):
                return "issue"
        if document_type == "report":
            if re.search(r"\bproblem\b|\bchallenge\b|\bpain point\b", section_text):
                return "problem"
            if re.search(r"\bcurrent state\b|\bas-is\b|\bbaseline\b", section_text):
                return "current_state"
            if re.search(r"\bgoal\b|\btarget\b|\bobjective\b", section_text):
                return "goal"
            if re.search(r"\brecommendation\b|\bproposal\b|\bsolution\b", section_text):
                return "recommendation"
        if document_type == "reference":
            if re.search(r"\bdefinition\b|\bmeaning\b|\brefers to\b", section_text):
                return "definition"
            if re.search(r"\bfield\b|\bparameter\b|\binput\b|\boutput\b", section_text):
                return "field_reference"
            if re.search(r"\bexample\b|\bsample\b", section_text):
                return "example"
        return "general"

    for chunk in chunks:
        metadata = dict(chunk.get("metadata") or {})
        metadata["sectionRole"] = infer_role(chunk)
        metadata["documentType"] = document_type
        chunk["metadata"] = metadata
    return chunks


def split_into_chunks(text: str, max_words: int | None = None, overlap: int | None = None) -> list[dict]:
    runtime = load_runtime_snapshot()
    max_words = runtime.chunk_size_words if max_words is None else max_words
    overlap = runtime.chunk_overlap_words if overlap is None else overlap
    if runtime.chunk_mode == "structured":
        structured_chunks = split_into_structured_chunks(promote_heading_lines(text), max_words=max_words, overlap=overlap)
        if structured_chunks:
            return structured_chunks
    return split_into_window_chunks(text, max_words=max_words, overlap=overlap)


def apply_semantic_unit_chunking(chunks: list[dict], document_type: str, max_words: int) -> list[dict]:
    semantic_chunks: list[dict] = []
    for chunk in chunks:
        content = str(chunk.get("content") or "").strip()
        metadata = dict(chunk.get("metadata") or {})
        role = str(metadata.get("sectionRole") or "general")
        if len(content.split()) <= max_words or role not in {"step", "rule", "exception", "decision", "recommendation", "definition", "scope"}:
            semantic_chunks.append(chunk)
            continue

        parts = [part.strip() for part in re.split(r"\n{2,}", content) if part.strip()]
        if len(parts) <= 1:
            semantic_chunks.append(chunk)
            continue

        for index, part in enumerate(parts, start=1):
            part_metadata = {
                **metadata,
                "chunkingMode": "semantic",
                "semanticUnit": role,
                "semanticPartIndex": index,
                "semanticPartCount": len(parts),
                "documentType": document_type,
            }
            semantic_chunks.append(
                {
                    "section_title": chunk.get("section_title"),
                    "content": part,
                    "token_count": len(tokenize(part)),
                    "metadata": part_metadata,
                }
            )
    return semantic_chunks


def build_section_summaries(chunks: list[dict], document_type: str) -> list[dict]:
    grouped: dict[str, list[dict]] = {}
    for chunk in chunks:
        metadata = chunk.get("metadata") or {}
        heading_path = metadata.get("headingPath") or []
        section_key = " > ".join([str(part) for part in heading_path if part]) or str(chunk.get("section_title") or "Document")
        grouped.setdefault(section_key, []).append(chunk)

    summaries: list[dict] = []
    for index, (section_key, section_chunks) in enumerate(grouped.items(), start=1):
        combined = "\n\n".join(str(chunk.get("content") or "").strip() for chunk in section_chunks if str(chunk.get("content") or "").strip())
        role_counts = Counter(str((chunk.get("metadata") or {}).get("sectionRole") or "general") for chunk in section_chunks)
        top_roles = [role for role, _count in role_counts.most_common(3)]
        sentence = SENTENCE_RE.split(combined.strip())[0].strip() if combined.strip() else ""
        summary = sentence[:280] if sentence else combined[:280]
        heading_path = list((section_chunks[0].get("metadata") or {}).get("headingPath") or [])
        summaries.append(
            {
                "sectionKey": f"sec-{slugify(section_key) or index}",
                "title": section_chunks[0].get("section_title") or section_key,
                "headingPath": heading_path,
                "summary": summary,
                "documentType": document_type,
                "chunkCount": len(section_chunks),
                "roles": top_roles,
                "firstChunkIndex": min(int(chunk.get("chunk_index_override", idx)) for idx, chunk in enumerate(section_chunks)),
            }
        )
    return summaries


def build_source_sections(chunks: list[dict], section_summaries: list[dict], document_type: str) -> list[dict]:
    grouped: dict[str, list[dict]] = {}
    for chunk in chunks:
        metadata = chunk.get("metadata") or {}
        heading_path = metadata.get("headingPath") or []
        section_key = " > ".join([str(part) for part in heading_path if part]) or str(chunk.get("section_title") or "Document")
        grouped.setdefault(section_key, []).append(chunk)

    sections: list[dict] = []
    summary_map = {str(item.get("sectionKey") or ""): item for item in section_summaries}
    summary_by_title = {str(item.get("title") or ""): item for item in section_summaries}
    for index, (section_key, section_chunks) in enumerate(grouped.items(), start=1):
        first = section_chunks[0]
        first_metadata = first.get("metadata") or {}
        heading_path = list(first_metadata.get("headingPath") or [])
        title = first.get("section_title") or section_key
        summary_key = f"sec-{slugify(section_key) or index}"
        summary_item = summary_map.get(summary_key) or summary_by_title.get(str(title))
        sections.append(
            {
                "sectionKey": summary_item.get("sectionKey") if summary_item else summary_key,
                "title": title,
                "headingPath": heading_path,
                "documentType": document_type,
                "roles": _unique_preserve_order([str((chunk.get("metadata") or {}).get("sectionRole") or "general") for chunk in section_chunks]),
                "chunkIndexes": [int((chunk.get("metadata") or {}).get("chunkIndex") or 0) for chunk in section_chunks],
                "pageRange": {
                    "start": min(int((chunk.get("metadata") or {}).get("pageNumber") or idx + 1) for idx, chunk in enumerate(section_chunks)),
                    "end": max(int((chunk.get("metadata") or {}).get("pageNumber") or idx + 1) for idx, chunk in enumerate(section_chunks)),
                },
                "summary": (summary_item or {}).get("summary") if summary_item else None,
                "bodyPreview": "\n\n".join(str(chunk.get("content") or "").strip() for chunk in section_chunks)[:600],
            }
        )
    return sections


def build_chunk_profile(chunks: list[dict], document_profile: dict, section_summaries: list[dict], source_sections: list[dict]) -> dict:
    role_counts = Counter(str((chunk.get("metadata") or {}).get("sectionRole") or "general") for chunk in chunks)
    chunk_mode_counts = Counter(str((chunk.get("metadata") or {}).get("chunkingMode") or "window") for chunk in chunks)
    semantic_unit_counts = Counter(str((chunk.get("metadata") or {}).get("semanticUnit") or "") for chunk in chunks if (chunk.get("metadata") or {}).get("semanticUnit"))
    avg_tokens = round(sum(int(chunk.get("token_count") or 0) for chunk in chunks) / len(chunks), 2) if chunks else 0
    return {
        "documentType": document_profile["documentType"],
        "documentTypeConfidence": document_profile["confidence"],
        "documentTypeReasons": list(document_profile.get("reasons") or []),
        "sourceType": document_profile.get("sourceType"),
        "chunkCount": len(chunks),
        "sectionCount": len(section_summaries),
        "sourceSectionCount": len(source_sections),
        "avgTokenCount": avg_tokens,
        "chunkModeCounts": dict(chunk_mode_counts),
        "roleCounts": dict(role_counts),
        "semanticUnitCounts": dict(semantic_unit_counts),
    }


def split_into_window_chunks(text: str, max_words: int, overlap: int) -> list[dict]:
    paragraphs = [paragraph.strip() for paragraph in text.split("\n\n") if paragraph.strip()]
    chunks: list[dict] = []
    current: list[str] = []
    current_words = 0
    section_title = "Document"

    def flush_chunk() -> None:
        nonlocal current, current_words
        if not current:
            return
        content = "\n\n".join(current).strip()
        chunks.append({"section_title": section_title, "content": content, "token_count": len(tokenize(content))})
        if overlap > 0:
            words = content.split()
            current = [" ".join(words[-overlap:])] if len(words) > overlap else [content]
            current_words = len(current[0].split())
        else:
            current = []
            current_words = 0

    for paragraph in paragraphs:
        if paragraph.startswith("#"):
            section_title = paragraph.lstrip("#").strip() or section_title

        paragraph_words = len(paragraph.split())
        if current and current_words + paragraph_words > max_words:
            flush_chunk()
        current.append(paragraph)
        current_words += paragraph_words

    flush_chunk()
    if not chunks and text.strip():
        chunks.append({"section_title": section_title, "content": text.strip(), "token_count": len(tokenize(text))})
    return chunks


def split_into_structured_chunks(text: str, max_words: int, overlap: int) -> list[dict]:
    blocks = _extract_markdown_blocks(text)
    if not blocks:
        return []

    expanded_blocks: list[dict] = []
    for block in blocks:
        expanded_blocks.extend(_split_oversized_block(block, max_words=max_words, overlap=overlap))

    chunks: list[dict] = []
    pending_blocks: list[dict] = []
    pending_words = 0
    current_heading_path: list[str] = []

    def flush_pending() -> None:
        nonlocal pending_blocks, pending_words, current_heading_path
        if not pending_blocks:
            return
        content = "\n\n".join(block["content"].strip() for block in pending_blocks if block["content"].strip()).strip()
        if not content:
            pending_blocks = []
            pending_words = 0
            current_heading_path = []
            return
        block_types = _unique_preserve_order([str(block.get("type") or "paragraph") for block in pending_blocks])
        heading_path = list(current_heading_path)
        section_title = heading_path[-1] if heading_path else str(pending_blocks[0].get("section_title") or "Document")
        chunks.append(
            {
                "section_title": section_title,
                "content": content,
                "token_count": len(tokenize(content)),
                "metadata": {
                    "chunkingMode": "structured",
                    "blockTypes": block_types,
                    "headingPath": heading_path,
                    "blockCount": len(pending_blocks),
                },
            }
        )
        pending_blocks = []
        pending_words = 0
        current_heading_path = []

    for block in expanded_blocks:
        block_type = str(block.get("type") or "paragraph")
        block_words = len(str(block.get("content") or "").split())
        heading_path = list(block.get("heading_path") or [])
        is_standalone = block_type in {"table", "list", "code"}

        if pending_blocks and (heading_path != current_heading_path or (not is_standalone and pending_words + block_words > max_words)):
            flush_pending()

        if is_standalone:
            flush_pending()
            content = str(block.get("content") or "").strip()
            if not content:
                continue
            chunks.append(
                {
                    "section_title": str(block.get("section_title") or (heading_path[-1] if heading_path else "Document")),
                    "content": content,
                    "token_count": len(tokenize(content)),
                    "metadata": {
                        "chunkingMode": "structured",
                        "blockTypes": [block_type],
                        "headingPath": heading_path,
                        "blockCount": 1,
                    },
                }
            )
            continue

        pending_blocks.append(block)
        pending_words += block_words
        current_heading_path = heading_path

    flush_pending()
    return chunks


def _extract_markdown_blocks(text: str) -> list[dict]:
    lines = text.splitlines()
    blocks: list[dict] = []
    heading_stack: list[str] = []
    paragraph_lines: list[str] = []
    list_lines: list[str] = []
    table_lines: list[str] = []
    code_lines: list[str] = []
    in_code_block = False

    def current_section_title() -> str:
        return heading_stack[-1] if heading_stack else "Document"

    def heading_path() -> list[str]:
        return list(heading_stack)

    def flush_paragraph() -> None:
        nonlocal paragraph_lines
        content = normalize_text("\n".join(paragraph_lines))
        if content:
            blocks.append({"type": "paragraph", "content": content, "section_title": current_section_title(), "heading_path": heading_path()})
        paragraph_lines = []

    def flush_list() -> None:
        nonlocal list_lines
        content = normalize_text("\n".join(list_lines))
        if content:
            blocks.append({"type": "list", "content": content, "section_title": current_section_title(), "heading_path": heading_path()})
        list_lines = []

    def flush_table() -> None:
        nonlocal table_lines
        content = normalize_text("\n".join(table_lines))
        if content:
            blocks.append({"type": "table", "content": content, "section_title": current_section_title(), "heading_path": heading_path()})
        table_lines = []

    def flush_code() -> None:
        nonlocal code_lines
        content = "\n".join(code_lines).strip()
        if content:
            blocks.append({"type": "code", "content": content, "section_title": current_section_title(), "heading_path": heading_path()})
        code_lines = []

    def flush_all() -> None:
        flush_paragraph()
        flush_list()
        flush_table()
        flush_code()

    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            flush_paragraph()
            flush_list()
            flush_table()
            code_lines.append(line)
            if in_code_block:
                in_code_block = False
                flush_code()
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            flush_all()
            level = len(heading_match.group(1))
            title = normalize_text(heading_match.group(2))
            if title:
                heading_stack[:] = heading_stack[: level - 1]
                heading_stack.append(title)
            continue

        if not stripped:
            flush_paragraph()
            flush_list()
            flush_table()
            continue

        if re.match(r"^(\*|-|\+|\d+\.)\s+", stripped):
            flush_paragraph()
            flush_table()
            list_lines.append(stripped)
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            flush_paragraph()
            flush_list()
            table_lines.append(stripped)
            continue

        flush_list()
        flush_table()
        paragraph_lines.append(stripped)

    flush_all()
    return blocks


def _split_oversized_block(block: dict, max_words: int, overlap: int) -> list[dict]:
    content = str(block.get("content") or "").strip()
    if len(content.split()) <= max_words:
        return [block]
    if str(block.get("type") or "") in {"table", "code"}:
        return [block]

    sentences = [sentence.strip() for sentence in SENTENCE_RE.split(content) if sentence.strip()]
    if len(sentences) <= 1:
        return [block]

    split_blocks: list[dict] = []
    current: list[str] = []
    current_words = 0

    def flush_sentence_chunk() -> None:
        nonlocal current, current_words
        if not current:
            return
        chunk_content = " ".join(current).strip()
        split_blocks.append({**block, "content": chunk_content})
        if overlap > 0:
            words = chunk_content.split()
            carry = words[-overlap:] if len(words) > overlap else words
            current = [" ".join(carry)] if carry else []
            current_words = len(carry)
        else:
            current = []
            current_words = 0

    for sentence in sentences:
        sentence_words = len(sentence.split())
        if current and current_words + sentence_words > max_words:
            flush_sentence_chunk()
        current.append(sentence)
        current_words += sentence_words
    flush_sentence_chunk()
    return split_blocks or [block]


def _unique_preserve_order(items: list[str]) -> list[str]:
    seen: set[str] = set()
    values: list[str] = []
    for item in items:
        if item in seen:
            continue
        seen.add(item)
        values.append(item)
    return values


def parse_docx_blocks(document: DocxDocument, path: Path) -> tuple[list[dict], list[str]]:
    upload_dir = ensure_upload_dir()
    assets_dir = upload_dir / f"{path.stem}-assets"
    if assets_dir.exists():
        shutil.rmtree(assets_dir)
    assets_dir.mkdir(parents=True, exist_ok=True)

    relationship_urls: dict[str, str] = {}
    for index, (relationship_id, relationship) in enumerate(sorted(document.part.rels.items()), start=1):
        if "image" not in relationship.reltype:
            continue
        suffix = Path(getattr(relationship.target_part, "partname", "")).suffix.lower() or ".bin"
        target_path = assets_dir / f"image-{index}{suffix}"
        target_path.write_bytes(relationship.target_part.blob)
        relationship_urls[relationship_id] = public_upload_url(target_path)

    ordered_blocks: list[dict] = []
    for block in iter_docx_blocks(document):
        if isinstance(block, Paragraph):
            text = normalize_text(block.text)
            style_name = block.style.name if block.style is not None else ""
            if text:
                ordered_blocks.append({"type": "paragraph", "content": text, "style": style_name})
            for image_url in extract_paragraph_image_urls(block, relationship_urls):
                ordered_blocks.append({"type": "image", "url": image_url, "alt": style_name or "Illustration"})
        elif isinstance(block, Table):
            table_markdown = table_to_markdown(block)
            if table_markdown:
                ordered_blocks.append({"type": "table", "content": table_markdown, "style": "Table"})

    image_urls = [block["url"] for block in ordered_blocks if block["type"] == "image"]
    if not relationship_urls:
        shutil.rmtree(assets_dir, ignore_errors=True)
    return ordered_blocks, image_urls


def iter_docx_blocks(document: DocxDocument):
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def extract_paragraph_image_urls(paragraph: Paragraph, relationship_urls: dict[str, str]) -> list[str]:
    image_urls: list[str] = []
    for drawing in paragraph._p.xpath(DRAWING_XPATH):
        relationship_id = drawing.get(qn("r:embed"))
        image_url = relationship_urls.get(relationship_id)
        if image_url:
            image_urls.append(image_url)
    return image_urls


def table_to_markdown(table: Table) -> str:
    rows = []
    for row in table.rows:
        cells = [" ".join(cell.text.split()) for cell in row.cells]
        if any(cells):
            rows.append(cells)
    if not rows:
        return ""
    column_count = max(len(row) for row in rows)
    normalized_rows = [row + [""] * (column_count - len(row)) for row in rows]
    header = normalized_rows[0]
    separator = ["---"] * column_count
    lines = ["| " + " | ".join(header) + " |", "| " + " | ".join(separator) + " |"]
    for row in normalized_rows[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def public_upload_url(path: Path) -> str:
    relative = path.relative_to(ensure_upload_dir()).as_posix()
    return f"/backend-uploads/{quote(relative)}"


def summarize_text(title: str, text: str) -> tuple[str, list[str]]:
    sentences = [sentence.strip() for sentence in SENTENCE_RE.split(text) if len(sentence.strip()) > 30]
    summary_sentences = sentences[:2] or [text[:280].strip()]
    key_facts = summary_sentences[:3]
    heuristic_summary = " ".join(summary_sentences)[:500].strip()

    runtime = load_runtime_snapshot()
    ingest_profile = runtime.profile_for_task("ingest_summary")
    if not llm_client.is_enabled(ingest_profile):
        return heuristic_summary, key_facts

    system_prompt = (
        "You summarize internal knowledge base source material. "
        "Return strict JSON with keys summary and key_facts. "
        "summary must be one concise paragraph. key_facts must be an array of up to 4 factual bullets."
    )
    user_prompt = f"Title: {title}\n\nContent:\n{text[:12000]}"
    response = llm_client.complete(ingest_profile, system_prompt, user_prompt)
    if not response:
        return heuristic_summary, key_facts

    try:
        payload = json_like_to_dict(response)
        summary = str(payload.get("summary") or heuristic_summary).strip()
        facts = [str(item).strip() for item in payload.get("key_facts", []) if str(item).strip()]
        return summary, facts[:4] or key_facts
    except Exception:
        return heuristic_summary, key_facts


def prompt_metadata(stage: str, provider: str, model: str) -> dict:
    return {
        "promptVersion": PROMPT_VERSION,
        "stage": stage,
        "provider": provider or "none",
        "model": model or "",
    }


def json_like_to_dict(raw: str) -> dict:
    start = raw.find("{")
    end = raw.rfind("}")
    if start == -1 or end == -1:
        raise ValueError("JSON block not found")
    import json

    return json.loads(raw[start : end + 1])


def extract_entities(text: str, limit: int = 8) -> list[dict]:
    candidates = re.findall(r"\b(?:[A-Z][a-zA-Z0-9]+|[A-Z]{2,})(?:\s+(?:[A-Z][a-zA-Z0-9]+|[A-Z]{2,})){0,3}\b", text)
    stopwords = {"The", "This", "That", "These", "Those", "Document", "Section", "Chapter", "Figure", "Table", "Overview", "Summary", "Source"}
    counts: Counter[str] = Counter()
    for candidate in candidates:
        normalized_candidate = " ".join(candidate.split()).strip()
        if len(normalized_candidate) <= 3 or normalized_candidate in stopwords:
            continue
        if normalized_candidate.isupper() and len(normalized_candidate) > 8:
            continue
        counts[normalized_candidate] += len(re.findall(rf"\b{re.escape(normalized_candidate)}\b", text))

    entities = []
    for name, mention_count in counts.most_common(limit):
        normalized = slugify(name).replace("-", " ")
        entities.append(
            {
                "id": f"ent-gen-{slugify(name)}",
                "name": name,
                "entity_type": infer_entity_type(name),
                "description": f"Entity extracted from source content: {name}",
                "aliases": [],
                "normalized_name": normalized,
                "mention_count": mention_count,
                "confidence_score": round(min(0.55 + (mention_count * 0.08), 0.95), 2),
            }
        )
    return entities


def infer_entity_type(name: str) -> str:
    lowered = name.lower()
    if any(word in lowered for word in ["framework", "process", "workflow", "pipeline"]):
        return "process"
    if any(word in lowered for word in ["ai", "llm", "model", "database", "api", "rag", "system"]):
        return "technology"
    if any(word in lowered for word in ["inc", "corp", "company", "organization"]):
        return "organization"
    return "concept"


CLAIM_TYPES = {"fact", "definition", "process", "rule", "example", "inference", "requirement", "condition", "decision", "metric", "risk", "instruction"}
VAGUE_TERMS = {"can", "could", "may", "might", "various", "several", "sometimes", "generally", "typically", "often"}


def _claim_span_from_excerpt(content: str, excerpt: str) -> tuple[int | None, int | None, str]:
    body = content or ""
    candidate = (excerpt or "").strip()
    if not body or not candidate:
        return None, None, ""
    direct = body.find(candidate)
    if direct != -1:
        return direct, direct + len(candidate), candidate
    normalized_candidate = " ".join(candidate.split())
    normalized_body = " ".join(body.split())
    normalized_index = normalized_body.find(normalized_candidate)
    if normalized_index != -1:
        start_term = normalized_candidate.split()[0]
        fallback = body.lower().find(start_term.lower())
        if fallback != -1:
            end = min(len(body), fallback + len(candidate))
            return fallback, end, body[fallback:end].strip()
    return None, None, ""


def _claim_type_from_text(sentence: str, topic: str | None = None) -> str:
    lowered = sentence.lower()
    topic_lowered = (topic or "").lower()
    if re.search(r"\bif\b|\bwhen\b|\bonce\b|\bunless\b|\bprovided that\b", lowered):
        return "condition"
    if re.search(r"\bmust\b|\brequired\b|\bshall\b|\bneed to\b|\bmandatory\b", lowered):
        return "requirement"
    if re.search(r"\bshould\b|\bpolicy\b|\brule\b|\bprohibited\b|\bnot allowed\b|\bforbidden\b", lowered):
        return "rule"
    if re.search(r"\bapprove\b|\breject\b|\bdecision\b|\bescalat", lowered):
        return "decision"
    if re.search(r"\brisk\b|\bimpact\b|\bfailure\b|\bincident\b|\berror\b|\bwarning\b", lowered):
        return "risk"
    if re.search(r"\bstep\b|\bworkflow\b|\bprocess\b|\bprocedure\b|\bsubmit\b|\breview\b|\bpublish\b|\bvalidate\b", lowered) or "sop" in topic_lowered:
        return "process"
    if re.search(r"\bmeans\b|\brefers to\b|\bis defined as\b|\bare defined as\b", lowered):
        return "definition"
    if re.search(r"\b\d+(?:\.\d+)?%?\b", lowered) and re.search(r"\bthreshold\b|\btarget\b|\bsla\b|\baccuracy\b|\blatency\b|\bscore\b|\bpercent\b|\bhour\b|\bday\b", lowered):
        return "metric"
    if re.search(r"\bfor example\b|\be.g\.\b|\bsuch as\b", lowered):
        return "example"
    if re.search(r"\bclick\b|\bopen\b|\bselect\b|\benter\b|\bconfigure\b|\brun\b", lowered):
        return "instruction"
    return "fact"


def _claim_confidence(sentence: str, claim_type: str, topic: str | None = None) -> tuple[float, list[str]]:
    lowered = sentence.lower()
    reasons: list[str] = []
    score = 0.58
    if len(sentence.split()) >= 8:
        score += 0.08
        reasons.append("sentence_length")
    if claim_type in {"requirement", "rule", "metric", "process", "decision", "instruction"}:
        score += 0.12
        reasons.append("typed_signal")
    if re.search(r"\b\d+(?:\.\d+)?%?\b", lowered):
        score += 0.08
        reasons.append("numeric_anchor")
    if topic and topic != "Document":
        score += 0.04
        reasons.append("section_topic")
    vague_hits = [term for term in VAGUE_TERMS if re.search(rf"\b{re.escape(term)}\b", lowered)]
    if vague_hits:
        score -= 0.16
        reasons.append("vague_language")
    if len(sentence.split()) < 6:
        score -= 0.1
        reasons.append("too_short")
    return max(0.25, min(score, 0.96)), vague_hits


def _claim_topic(chunk: dict) -> str:
    metadata = chunk.get("metadata") or {}
    heading_path = metadata.get("headingPath") or []
    if heading_path:
        return str(heading_path[-1])[:120]
    return str(chunk.get("section_title") or "Document")[:120]


def _heuristic_claims_from_chunk(chunk: dict, chunk_index: int, entity_ids: dict[str, str]) -> list[dict]:
    claims: list[dict] = []
    content = chunk["content"]
    topic = _claim_topic(chunk)
    sentences = [sentence.strip() for sentence in SENTENCE_RE.split(content) if len(sentence.strip()) > 25]
    for sentence in sentences[:3]:
        claim_type = _claim_type_from_text(sentence, topic=topic)
        confidence, vague_hits = _claim_confidence(sentence, claim_type, topic=topic)
        span_start, span_end, evidence_excerpt = _claim_span_from_excerpt(content, sentence)
        if span_start is None or span_end is None:
            continue
        referenced_entities = [entity_id for name, entity_id in entity_ids.items() if name.lower() in sentence.lower()]
        metadata_json = {
            "grounding": "sentence",
            "vagueTerms": vague_hits,
            "isLowConfidence": confidence < 0.62,
            "headingPath": (chunk.get("metadata") or {}).get("headingPath", []),
            "blockTypes": (chunk.get("metadata") or {}).get("blockTypes", []),
            "promptVersion": PROMPT_VERSION,
        }
        if metadata_json["isLowConfidence"]:
            metadata_json["reviewHint"] = "low_confidence_claim"
        claims.append(
            {
                "id": f"clm-{uuid4().hex[:8]}",
                "text": sentence[:500],
                "claim_type": claim_type,
                "confidence_score": round(confidence, 2),
                "canonical_status": "unverified",
                "review_status": "pending",
                "topic": topic,
                "entity_ids": referenced_entities,
                "chunk_index": chunk_index,
                "extraction_method": "heuristic",
                "evidence_span_start": span_start,
                "evidence_span_end": span_end,
                "metadata_json": {**metadata_json, "evidenceExcerpt": evidence_excerpt[:280]},
            }
        )
    return claims


def _normalize_llm_claim_item(item: dict, chunk: dict, chunk_index: int, entity_ids: dict[str, str]) -> dict | None:
    text = " ".join(str(item.get("text") or item.get("statement") or "").split()).strip()
    if len(text) < 20:
        return None
    claim_type = str(item.get("claim_type") or item.get("type") or "").strip().lower() or _claim_type_from_text(text, _claim_topic(chunk))
    if claim_type not in CLAIM_TYPES:
        claim_type = _claim_type_from_text(text, _claim_topic(chunk))
    raw_confidence = item.get("confidence", item.get("confidence_score", 0.72))
    try:
        confidence = float(raw_confidence)
    except Exception:
        confidence = 0.72
    evidence_excerpt = " ".join(str(item.get("evidence_excerpt") or item.get("evidence") or text).split()).strip()
    span_start, span_end, evidence_excerpt = _claim_span_from_excerpt(chunk["content"], evidence_excerpt)
    if span_start is None or span_end is None:
        span_start, span_end, evidence_excerpt = _claim_span_from_excerpt(chunk["content"], text)
    if span_start is None or span_end is None:
        return None
    mentioned_entities = item.get("entities", [])
    referenced_entities = []
    if isinstance(mentioned_entities, list):
        for entity in mentioned_entities:
            name = str(entity.get("name") if isinstance(entity, dict) else entity).strip()
            if name and name in entity_ids:
                referenced_entities.append(entity_ids[name])
            elif name:
                matched = next((entity_id for label, entity_id in entity_ids.items() if label.lower() == name.lower()), None)
                if matched:
                    referenced_entities.append(matched)
    if not referenced_entities:
        referenced_entities = [entity_id for name, entity_id in entity_ids.items() if name.lower() in text.lower()]
    vague_hits = [term for term in VAGUE_TERMS if re.search(rf"\b{re.escape(term)}\b", text.lower())]
    metadata_json = {
        "grounding": "llm_json",
        "evidenceExcerpt": evidence_excerpt[:280],
        "rawType": item.get("claim_type") or item.get("type"),
        "reviewHint": "low_confidence_claim" if confidence < 0.62 else "",
        "headingPath": (chunk.get("metadata") or {}).get("headingPath", []),
        "blockTypes": (chunk.get("metadata") or {}).get("blockTypes", []),
        "vagueTerms": vague_hits,
        "isLowConfidence": confidence < 0.62,
        "promptVersion": PROMPT_VERSION,
    }
    topic = _claim_topic(chunk)
    return {
        "id": f"clm-{uuid4().hex[:8]}",
        "text": text[:500],
        "claim_type": claim_type,
        "confidence_score": round(max(0.25, min(confidence, 0.99)), 2),
        "canonical_status": "unverified",
        "review_status": "pending",
        "topic": topic,
        "entity_ids": _unique_preserve_order(referenced_entities),
        "chunk_index": chunk_index,
        "extraction_method": "llm",
        "evidence_span_start": span_start,
        "evidence_span_end": span_end,
        "metadata_json": metadata_json,
    }


def _llm_claims_from_chunk(chunk: dict, chunk_index: int, entity_ids: dict[str, str]) -> list[dict]:
    runtime = load_runtime_snapshot()
    claim_profile = runtime.profile_for_task("claim_extraction")
    if not llm_client.is_enabled(claim_profile):
        return []
    content = chunk["content"]
    topic = _claim_topic(chunk)
    system_prompt = (
        "You extract source-grounded claims from a document chunk. "
        "Return strict JSON only. "
        "Do not invent unsupported claims. "
        "Each claim must be atomic and grounded in the provided chunk. "
        "Use claim_type from: fact, definition, process, rule, example, inference, requirement, condition, decision, metric, risk, instruction. "
        "Include evidence_excerpt copied from the chunk."
    )
    user_prompt = (
        f"Section: {topic}\n"
        f"Heading path: {(chunk.get('metadata') or {}).get('headingPath', [])}\n"
        "Return JSON object with key claims.\n\n"
        f"Chunk:\n{content[:6000]}"
    )
    response = llm_client.complete(claim_profile, system_prompt, user_prompt)
    if not response:
        return []
    try:
        payload = json_like_to_dict(response)
    except Exception:
        return []
    claims: list[dict] = []
    for item in payload.get("claims", [])[:4]:
        if not isinstance(item, dict):
            continue
        normalized = _normalize_llm_claim_item(item, chunk, chunk_index, entity_ids)
        if normalized:
            claims.append(normalized)
    return claims


def build_claims(chunks: list[dict], entity_ids: dict[str, str]) -> list[dict]:
    claims: list[dict] = []
    seen_texts: set[str] = set()
    for index, chunk in enumerate(chunks[:12]):
        llm_claims = _llm_claims_from_chunk(chunk, index, entity_ids)
        heuristic_claims = _heuristic_claims_from_chunk(chunk, index, entity_ids)
        if llm_claims:
            claims.extend(llm_claims)
            existing_texts = {item["text"].lower() for item in llm_claims}
            claims.extend([item for item in heuristic_claims if item["text"].lower() not in existing_texts][:1])
        else:
            claims.extend(heuristic_claims)
    validated: list[dict] = []
    for claim in claims:
        normalized = " ".join(str(claim.get("text") or "").lower().split())
        if not normalized or normalized in seen_texts:
            continue
        seen_texts.add(normalized)
        if len(normalized.split()) > 80:
            continue
        if claim.get("evidence_span_start") is None or claim.get("evidence_span_end") is None:
            continue
        validated.append(claim)
    return validated


def extract_timeline_events(text: str, entities: list[dict], limit: int = 12) -> list[dict]:
    entity_lookup = {entity["name"].lower(): entity["id"] for entity in entities}
    events: list[dict] = []
    seen: set[tuple[str, str]] = set()
    sentences = [sentence.strip() for sentence in SENTENCE_RE.split(text) if len(sentence.strip()) > 20]

    def append_event(date_value: str, sort_key: str, precision: str, sentence: str) -> None:
        title = sentence[:120].rstrip(".")
        key = (sort_key, title.lower())
        if key in seen:
            return
        seen.add(key)
        linked_entities = [entity_id for name, entity_id in entity_lookup.items() if name in sentence.lower()]
        events.append(
            {
                "id": f"evt-{uuid4().hex[:8]}",
                "event_date": date_value,
                "sort_key": sort_key,
                "precision": precision,
                "title": title,
                "description": sentence[:320],
                "entity_ids": linked_entities,
            }
        )

    for sentence in sentences:
        for match in re.finditer(r"\b(20\d{2})-(\d{2})-(\d{2})\b", sentence):
            year, month, day = match.groups()
            append_event(match.group(0), f"{year}-{month}-{day}", "day", sentence)
        for match in re.finditer(r"\b(20\d{2})/(\d{2})/(\d{2})\b", sentence):
            year, month, day = match.groups()
            append_event(match.group(0), f"{year}-{month}-{day}", "day", sentence)
        for match in re.finditer(r"\bQ([1-4])\s*(20\d{2})\b", sentence, flags=re.IGNORECASE):
            quarter, year = match.groups()
            month = f"{(int(quarter) - 1) * 3 + 1:02d}"
            append_event(f"Q{quarter} {year}", f"{year}-{month}-01", "quarter", sentence)
        for match in re.finditer(r"\b(20\d{2})\b", sentence):
            year = match.group(1)
            append_event(year, f"{year}-01-01", "year", sentence)

    return sorted(events, key=lambda item: item["sort_key"])[:limit]


def extract_glossary_terms(text: str, limit: int = 20) -> list[dict]:
    terms: list[dict] = []
    seen: set[str] = set()
    lines = [line.strip(" -\t") for line in text.splitlines() if line.strip()]
    patterns = [
        re.compile(r"^([A-Za-z][A-Za-z0-9 /()_-]{1,60}):\s+(.{12,400})$"),
        re.compile(r"^([A-Za-z][A-Za-z0-9 /()_-]{1,60})\s+-\s+(.{12,400})$"),
    ]

    for line in lines:
        for pattern in patterns:
            match = pattern.match(line)
            if not match:
                continue
            term = " ".join(match.group(1).split())
            definition = " ".join(match.group(2).split())
            normalized = slugify(term).replace("-", " ")
            if not normalized or normalized in seen:
                continue
            if len(term.split()) > 8 or len(definition.split()) < 3:
                continue
            seen.add(normalized)
            terms.append(
                {
                    "id": f"glo-{uuid4().hex[:8]}",
                    "term": term,
                    "normalized_term": normalized,
                    "definition": definition[:500],
                    "aliases": [],
                    "confidence_score": 0.76,
                }
            )
            break
        if len(terms) >= limit:
            break
    return terms


def classify_page_type_candidates(
    title: str,
    parsed: ParsedDocument,
    chunks: list[dict],
    entities: list[dict],
    timeline_events: list[dict],
    glossary_terms: list[dict],
) -> list[dict]:
    text = parsed.text
    lowered = f"{title}\n{text}".lower()
    candidates: list[dict] = [{"pageType": "summary", "confidence": 0.98, "reason": "Every source can generate a summary page."}]

    sop_score = 0.0
    if parsed.metadata.get("imageCount", 0) > 0:
        sop_score += 0.2
    if re.search(r"\bstep\b|\bworkflow\b|\bprocedure\b|\bchecklist\b", lowered):
        sop_score += 0.45
    if re.search(r"(?:^|\s)1\.(?:\s|$)|(?:^|\s)2\.(?:\s|$)|step\s+1|step\s+2", lowered):
        sop_score += 0.2
    if sop_score >= 0.4:
        candidates.append({"pageType": "sop", "confidence": round(min(sop_score, 0.95), 2), "reason": "Detected procedural language, ordered steps, or UI walkthrough cues."})

    concept_score = 0.0
    if re.search(r"\bis\b|\bare\b|\bdefinition\b|\bconcept\b|\boverview\b", lowered):
        concept_score += 0.35
    if len(chunks) <= 4:
        concept_score += 0.1
    if concept_score >= 0.3:
        candidates.append({"pageType": "concept", "confidence": round(min(concept_score, 0.8), 2), "reason": "Detected definition-style or explanatory content."})

    issue_score = 0.0
    if re.search(r"\bissue\b|\brisk\b|\bproblem\b|\bconflict\b|\bwarning\b|\bincident\b", lowered):
        issue_score += 0.55
    if issue_score >= 0.4:
        candidates.append({"pageType": "issue", "confidence": round(min(issue_score, 0.85), 2), "reason": "Detected issue, risk, or conflict vocabulary."})

    timeline_score = 0.0
    if timeline_events:
        timeline_score += min(0.2 + (len(timeline_events) * 0.08), 0.55)
    if re.search(r"\b20\d{2}\b|\bq[1-4]\b|\btimeline\b|\bmilestone\b|\broadmap\b", lowered):
        timeline_score += 0.45
    if timeline_score >= 0.4:
        candidates.append({"pageType": "timeline", "confidence": round(min(timeline_score, 0.8), 2), "reason": "Detected date/time-oriented language."})

    glossary_score = 0.0
    if glossary_terms:
        glossary_score += min(0.2 + (len(glossary_terms) * 0.06), 0.45)
    if text.count(":") >= 4:
        glossary_score += 0.35
    if re.search(r"\bterm\b|\bglossary\b|\bdefinition\b|\bacronym\b", lowered):
        glossary_score += 0.35
    if glossary_score >= 0.4:
        candidates.append({"pageType": "glossary", "confidence": round(min(glossary_score, 0.75), 2), "reason": "Detected repeated term-definition patterns."})

    entity_score = 0.0
    if len(tokenize(title)) >= 3:
        entity_score += 0.15
    if len(entities) >= 3:
        entity_score += 0.15
    if re.search(r"\bcompany\b|\bcustomer\b|\bsystem\b|\bproduct\b|\bproject\b|\bframework\b", lowered):
        entity_score += 0.25
    if entity_score >= 0.3:
        candidates.append({"pageType": "entity", "confidence": round(min(entity_score, 0.7), 2), "reason": "Detected entity-centric language in title or source body."})

    seen: set[str] = set()
    unique_candidates = []
    for candidate in candidates:
        if candidate["pageType"] in seen:
            continue
        seen.add(candidate["pageType"])
        unique_candidates.append(candidate)
    return unique_candidates


def run_ingest_pipeline(path: Path, mime_type: str, source_type: str, title: str) -> IngestArtifacts:
    parsed = parse_document(path, mime_type, source_type)
    chunks = split_into_chunks(parsed.text)
    summary, key_facts = summarize_text(title, parsed.text[:16000])
    document_profile = classify_document_profile(title, parsed, chunks, summary)
    parsed.metadata = {
        **(parsed.metadata or {}),
        "documentType": document_profile["documentType"],
        "documentTypeConfidence": document_profile["confidence"],
        "documentTypeReasons": document_profile["reasons"],
        "language": detect_language(parsed.text),
    }
    chunks = annotate_chunk_section_roles(chunks, document_profile["documentType"])
    runtime = load_runtime_snapshot()
    chunks = apply_semantic_unit_chunking(chunks, document_profile["documentType"], runtime.chunk_size_words)
    for chunk_index, chunk in enumerate(chunks):
        metadata = dict(chunk.get("metadata") or {})
        metadata["chunkIndex"] = chunk_index
        chunk["metadata"] = metadata
    section_summaries = build_section_summaries(chunks, document_profile["documentType"])
    source_sections = build_source_sections(chunks, section_summaries, document_profile["documentType"])
    parsed.metadata["sectionSummaries"] = section_summaries
    parsed.metadata["sourceSections"] = source_sections
    parsed.metadata["chunkProfile"] = build_chunk_profile(chunks, document_profile, section_summaries, source_sections)
    tags = build_tags(title, parsed.text)
    parsed.metadata["keywords"] = tags
    entities = extract_entities(parsed.text)
    entity_id_map = {entity["name"]: entity["id"] for entity in entities}
    claims = build_claims(chunks, entity_id_map)
    timeline_events = extract_timeline_events(parsed.text, entities)
    glossary_terms = extract_glossary_terms(parsed.text)
    page_type_candidates = classify_page_type_candidates(title, parsed, chunks, entities, timeline_events, glossary_terms)

    stage_results = [
        IngestStageResult(name="parse", status="completed", details={"sourceType": parsed.source_type, "charCount": len(parsed.text), "metadata": parsed.metadata}),
        IngestStageResult(
            name="chunk",
            status="completed",
            details={"chunkCount": len(chunks), "avgTokenCount": round(sum(chunk["token_count"] for chunk in chunks) / len(chunks), 2) if chunks else 0},
        ),
        IngestStageResult(name="summarize", status="completed", details={"summaryLength": len(summary), "keyFactCount": len(key_facts)}),
        IngestStageResult(name="extract_entities", status="completed", details={"entityCount": len(entities), "topEntities": [entity["name"] for entity in entities[:5]]}),
        IngestStageResult(name="extract_claims", status="completed", details={"claimCount": len(claims)}),
        IngestStageResult(name="extract_timeline", status="completed", details={"eventCount": len(timeline_events)}),
        IngestStageResult(name="extract_glossary", status="completed", details={"termCount": len(glossary_terms)}),
        IngestStageResult(name="classify_page_types", status="completed", details={"pageTypes": [candidate["pageType"] for candidate in page_type_candidates]}),
        IngestStageResult(name="classify_document_type", status="completed", details=document_profile),
        IngestStageResult(
            name="detect_section_roles",
            status="completed",
            details={
                "roles": {
                    role: count
                    for role, count in Counter(str((chunk.get("metadata") or {}).get("sectionRole") or "general") for chunk in chunks).items()
                }
            },
        ),
        IngestStageResult(
            name="build_section_summaries",
            status="completed",
            details={"sectionCount": len(section_summaries), "sectionTitles": [item["title"] for item in section_summaries[:8]]},
        ),
    ]

    return IngestArtifacts(
        parsed=parsed,
        chunks=chunks,
        summary=summary,
        key_facts=key_facts,
        tags=tags,
        entities=entities,
        claims=claims,
        timeline_events=timeline_events,
        glossary_terms=glossary_terms,
        page_type_candidates=page_type_candidates,
        stage_results=stage_results,
    )


def serialize_stage_results(stage_results: list[IngestStageResult]) -> list[dict]:
    return [asdict(result) for result in stage_results]


def build_page_markdown(
    title: str,
    summary: str,
    chunks: list[dict],
    key_facts: list[str],
    page_type: str = "summary",
    entities: list[dict] | None = None,
    timeline_events: list[dict] | None = None,
    glossary_terms: list[dict] | None = None,
    image_urls: list[str] | None = None,
    ordered_blocks: list[dict] | None = None,
    citation_markers: list[str] | None = None,
    citation_notes: list[str] | None = None,
) -> str:
    entities = entities or []
    timeline_events = timeline_events or []
    glossary_terms = glossary_terms or []
    lines = [f"# {title}", "", summary, ""]
    if key_facts:
        lines.extend(["## Key Facts", ""])
        markers = citation_markers or []
        lines.extend([f"- {fact}{markers[index] if index < len(markers) else ''}" for index, fact in enumerate(key_facts)])
        lines.append("")
    if page_type == "sop":
        lines.extend(["## Procedure", ""])
        step_chunks = chunks[:8] or [{"section_title": "Step 1", "content": summary}]
        for index, chunk in enumerate(step_chunks, start=1):
            lines.extend([f"### Step {index}: {chunk['section_title']}", "", chunk["content"], ""])
        lines.extend(["## Validation Checklist", ""])
        lines.extend(["- Confirm source evidence is complete.", "- Confirm owner/reviewer has approved the procedure.", "- Confirm screenshots or UI evidence match the latest workflow.", ""])
    elif page_type == "timeline" and timeline_events:
        lines.extend(["## Timeline", ""])
        for event in timeline_events[:20]:
            lines.extend([f"- **{event['event_date']}**: {event['title']}"])
        lines.append("")
        lines.extend(["## Notes", "", summary, ""])
    elif page_type == "glossary" and glossary_terms:
        lines.extend(["## Glossary Terms", ""])
        for term in glossary_terms[:30]:
            aliases = f" Alias: {', '.join(term.get('aliases') or [])}." if term.get("aliases") else ""
            lines.extend([f"- **{term['term']}**: {term['definition']}{aliases}"])
        lines.append("")
    elif page_type == "entity" and entities:
        primary = entities[0]
        lines.extend(["## Entity Profile", ""])
        lines.extend([f"- **Name:** {primary['name']}", f"- **Type:** {primary['entity_type']}", f"- **Description:** {primary['description']}", ""])
        if len(entities) > 1:
            lines.extend(["## Related Entities", ""])
            lines.extend([f"- {entity['name']} ({entity['entity_type']})" for entity in entities[1:8]])
            lines.append("")
    elif page_type == "issue":
        lines.extend(["## Issue Summary", "", summary, "", "## Risk And Impact", ""])
        lines.extend(["- Confirm affected systems, teams, or policies.", "- Capture impact, likelihood, and evidence before publishing.", ""])
        lines.extend(["## Evidence", ""])
        lines.extend([f"- {fact}" for fact in key_facts[:5]] or ["- Add source-backed evidence."])
        lines.append("")
    elif page_type == "concept":
        lines.extend(["## Concept", "", summary, "", "## How It Works", ""])
        lines.extend([chunk["content"] for chunk in chunks[:2]])
        lines.append("")
    if ordered_blocks:
        lines.extend(["## Source Walkthrough", ""])
        for index, block in enumerate(ordered_blocks, start=1):
            if block["type"] == "image":
                alt = block.get("alt") or f"Illustration {index}"
                lines.extend([f"![{alt}]({block['url']})", ""])
                continue
            if block["type"] == "table":
                lines.extend([block["content"], ""])
                continue
            content = block.get("content", "").strip()
            if not content:
                continue
            if block.get("style", "").startswith("Heading"):
                level_match = re.search(r"(\d+)", block["style"])
                level = min(max(int(level_match.group(1)) if level_match else 2, 2), 4)
                lines.extend([f"{'#' * level} {content}", ""])
            else:
                lines.extend([content, ""])
    elif image_urls:
        lines.extend(["## Illustrations", ""])
        for index, image_url in enumerate(image_urls, start=1):
            lines.extend([f"![Illustration {index}]({image_url})", ""])
    for chunk in chunks[:6]:
        lines.extend([f"## {chunk['section_title']}", "", chunk["content"], ""])
    if citation_notes:
        lines.extend(["## Source Citations", ""])
        lines.extend(citation_notes)
        lines.append("")
    return "\n".join(lines).strip()


def build_tags(title: str, text: str, limit: int = 6) -> list[str]:
    words = [token for token in tokenize(f"{title} {text}") if len(token) > 4]
    common = [word for word, _ in Counter(words).most_common(limit * 2)]
    unique = []
    for word in common:
        if word not in unique:
            unique.append(word)
        if len(unique) >= limit:
            break
    return unique


def score_text(query: str, text: str) -> float:
    query_terms = tokenize(query)
    if not query_terms:
        return 0.0
    text_terms = tokenize(text)
    if not text_terms:
        return 0.0
    text_counts = Counter(text_terms)
    overlap = sum(text_counts.get(term, 0) for term in query_terms)
    if overlap == 0:
        return 0.0
    return overlap / max(len(text_terms) ** 0.5, 1)


def slugify(value: str) -> str:
    return re.sub(r"[^a-z0-9-]", "", re.sub(r"\s+", "-", value.lower())).strip("-")
